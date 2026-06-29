"""Export campaign as JSON, ZIP, or DOCX."""
import json, io, zipfile, os, time

def export_json(brief: dict, results: dict, analytics: dict) -> bytes:
    data = {"brief": brief, "results": {k: v for k, v in results.items() if isinstance(v, str) and not os.path.isfile(str(v) or "")},
            "analytics": analytics, "exported_at": time.strftime("%Y-%m-%dT%H:%M:%S")}
    # include text results only (not file paths in json)
    data["results"] = {k: v for k, v in results.items() if k in ("tg", "bl", "sc")}
    return json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8")

def export_zip(brief: dict, results: dict, analytics: dict) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("campaign.json",  json.dumps({"brief": brief, "analytics": analytics}, indent=2))
        if results.get("tg"): z.writestr("tagline.txt", results["tg"])
        if results.get("bl"): z.writestr("blog.txt",    results["bl"])
        if results.get("sc"): z.writestr("social.json", json.dumps(results["sc"], indent=2, ensure_ascii=False))
        if results.get("im") and os.path.isfile(str(results["im"])):
            z.write(results["im"], "hero_image.png")
        if results.get("vd") and os.path.isfile(str(results["vd"])):
            z.write(results["vd"], "video.mp4")
        if analytics:
            z.writestr("analytics.json", json.dumps(analytics, indent=2))
    return buf.getvalue()

def export_docx(brief: dict, results: dict, analytics: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading("AI Marketing Campaign", 0)
        doc.add_heading("Campaign Brief", 1)
        for k, v in brief.items():
            doc.add_paragraph(f"{k.replace('_',' ').title()}: {v}")
        if results.get("tg"):
            doc.add_heading("Tagline", 1)
            doc.add_paragraph(results["tg"])
        if results.get("bl"):
            doc.add_heading("Blog Introduction", 1)
            doc.add_paragraph(results["bl"])
        if results.get("sc") and isinstance(results["sc"], dict):
            doc.add_heading("Social Media Posts", 1)
            for plat, txt in results["sc"].items():
                doc.add_heading(plat.title(), 2)
                doc.add_paragraph(str(txt))
        if analytics.get("scores"):
            doc.add_heading("Analytics", 1)
            for k, v in analytics["scores"].items():
                doc.add_paragraph(f"{k.replace('_',' ').title()}: {v}/100")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return export_json(brief, results, analytics)

def get_export_bytes(fmt: str, brief: dict, results: dict, analytics: dict) -> tuple:
    ts = time.strftime("%Y%m%d_%H%M")
    if fmt == "JSON":
        return export_json(brief, results, analytics), f"campaign_{ts}.json", "application/json"
    if fmt == "ZIP":
        return export_zip(brief, results, analytics),  f"campaign_{ts}.zip",  "application/zip"
    if fmt == "DOCX":
        return export_docx(brief, results, analytics), f"campaign_{ts}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return b"", "campaign.txt", "text/plain"
